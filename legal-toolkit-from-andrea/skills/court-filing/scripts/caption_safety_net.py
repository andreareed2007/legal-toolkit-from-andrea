#!/usr/bin/env python3
"""
Caption safety-net (multi-court).

If a court-type-specific validator's spec-conformance checks fail on a
docx-js-built filing, this script rebuilds a byte-correct caption with
python-docx and splices it into the target document.

Currently dispatches:
    --court-type tx-state  (TX state district courts)
    --court-type federal   (U.S. District Courts)
    --court-type aaa-arb   (AAA arbitration)

Other court types (business, ny-state) raise NotImplementedError until their
locked specs are written. AAA arbitration is implemented.

Usage:
    python3 caption_safety_net.py target.docx --court-type tx-state \\
        --case-input case_tx.json
    python3 caption_safety_net.py target.docx --court-type federal \\
        --case-input case_federal.json

case.json shape (TX state):
    {
      "cause_number": "DC-00-00000",
      "court_type_line": "DISTRICT COURT",
      "district": "44TH JUDICIAL DISTRICT",
      "county_line": "DALLAS COUNTY, TEXAS",
      "sides": [
        {"parties": [{"name": "ACME PARTNERS, INC.", "role": "Plaintiff"}]},
        {"parties": [{"name": "SAMPLE DEFENDANT", "role": "Defendant"}]}
      ]
    }

case.json shape (federal):
    {
      "case_number": "3:00-CV-00000-X",
      "case_number_prefix": "Civil Action No.",
      "district": "NORTHERN DISTRICT OF TEXAS",
      "division": "DALLAS DIVISION",
      "sides": [
        {"parties": [{"name": "ACME CORP, LLC", "role": "Plaintiff"}]},
        {"parties": [{"name": "SECOND DEFENDANT", "role": "Defendant"}]}
      ]
    }

case.json shape (AAA arbitration):
    {
      "case_number": "01-00-0000-0000",
      "sides": [
        {"parties": [{"name": "ALPHA INVESTMENT TRUST", "role": "Claimant"}]},
        {"parties": [{"name": "BETA INVESTMENT TRUST", "role": "Respondent"},
                     {"name": "GAMMA PE FUND 1, L.P. - SERIES 1", "role": "Respondent"}]}
      ]
    }

`sides` is the chain of party blocks separated by v.:
  - 2 sides       = P v. D                        (1 v., 9 rows)
  - 3 sides       = P v. D v. 3PD                 (2 v.'s, 15 rows)
  - N sides       = N-1 v.'s, 9 + 6 * (N-2) rows

Same-side parties go in the `parties` list -- they will be joined into a
single paragraph with Oxford comma + ALL CAPS "AND".
"""
import json
import sys
import argparse
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = 'Century Schoolbook'
SIZE = Pt(12)


# ─── Common helpers ──────────────────────────────────────────────────

def set_run_font(run, italic=False, bold=False):
    run.font.name = FONT
    run.font.size = SIZE
    run.font.italic = italic
    run.font.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rFonts.set(qn(attr), FONT)


def single_space(paragraph, space_after_pt=0):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(space_after_pt)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE


def add_para(container, text='', alignment=WD_ALIGN_PARAGRAPH.LEFT,
             italic=False, bold=False, left_indent=None,
             space_after_pt=0):
    p = container.add_paragraph()
    p.alignment = alignment
    single_space(p, space_after_pt=space_after_pt)
    if left_indent is not None:
        p.paragraph_format.left_indent = left_indent
    run = p.add_run(text if text else '')
    set_run_font(run, italic=italic, bold=bold)
    return p


def add_court_header(doc, lines):
    """Federal court header: single paragraph, line breaks, bold, centered,
    spaceAfter=12pt."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(12)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    for i, line in enumerate(lines):
        if i > 0:
            br_run = p.add_run()
            br_run.add_break()
            set_run_font(br_run, bold=True)
        run = p.add_run(line)
        set_run_font(run, bold=True)
    return p


def remove_table_borders(table):
    tblPr = table._element.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        table._element.insert(0, tblPr)
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement('w:' + side)
        b.set(qn('w:val'), 'nil')
        tblBorders.append(b)
    tblPr.append(tblBorders)


def remove_cell_borders(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    existing = tcPr.find(qn('w:tcBorders'))
    if existing is not None:
        tcPr.remove(existing)
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement('w:' + side)
        b.set(qn('w:val'), 'nil')
        tcBorders.append(b)
    tcPr.append(tcBorders)


def set_cell_margins_zero(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    existing = tcPr.find(qn('w:tcMar'))
    if existing is not None:
        tcPr.remove(existing)
    tcMar = OxmlElement('w:tcMar')
    for side in ('top', 'left', 'bottom', 'right'):
        m = OxmlElement('w:' + side)
        m.set(qn('w:w'), '0')
        m.set(qn('w:type'), 'dxa')
        tcMar.append(m)
    tcPr.append(tcMar)


def clear_cell(cell):
    for p in list(cell.paragraphs):
        p._element.getparent().remove(p._element)


def set_col_width(cell, width):
    tcPr = cell._tc.get_or_add_tcPr()
    existing = tcPr.find(qn('w:tcW'))
    if existing is not None:
        tcPr.remove(existing)
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(width.twips)))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)


# ─── Shared spec geometry ────────────────────────────────────────────

def _join_same_side_parties(parties):
    """Oxford comma + ALL CAPS AND."""
    names = [p["name"] for p in parties]
    if len(names) == 1:
        return names[0] + ","
    if len(names) == 2:
        return f"{names[0]} AND {names[1]},"
    return ", ".join(names[:-1]) + f", AND {names[-1]},"


def _role_label(parties, is_last_relationship):
    role = parties[0]["role"]
    plural = role + "s" if len(parties) > 1 and not role.endswith("s") else role
    terminator = "." if is_last_relationship else ","
    return f"{plural}{terminator}"


def _build_column_1_plan(sides):
    """Locked-spec layout for column 1, identical across TX state and federal."""
    out = []
    for i, side in enumerate(sides):
        is_last = (i == len(sides) - 1)
        if i == 0:
            out.append((_join_same_side_parties(side["parties"]), None))
            out.append(("", None))
            out.append((_role_label(side["parties"], is_last), 0.5))
        else:
            out.append(("", None))
            out.append(("v.", None))
            out.append(("", None))
            out.append((_join_same_side_parties(side["parties"]), None))
            out.append(("", None))
            out.append((_role_label(side["parties"], is_last), 0.5))
    return out


def _expected_row_count(sides):
    if len(sides) < 2:
        raise ValueError("Need at least 2 sides (one v.) for a caption.")
    return 3 + 6 * (len(sides) - 1)


def _first_v_row_index(sides):
    """0-indexed row of the first 'v.' in column 1 -- always row index 4 (row 5)."""
    return 4


def _common_doc_setup():
    """Return a Document with margins and Normal style configured."""
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    normal = doc.styles['Normal']
    normal.font.name = FONT
    normal.font.size = SIZE
    rpr = normal.element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rpr.append(rFonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rFonts.set(qn(attr), FONT)
    return doc


def _build_caption_table(doc):
    """Build the 1-row, 3-column caption table with locked widths/borders."""
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    remove_table_borders(table)

    widths = [Inches(3.0), Inches(0.5), Inches(3.0)]
    row = table.rows[0]
    for cell, width in zip(row.cells, widths):
        cell.width = width
        set_col_width(cell, width)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        remove_cell_borders(cell)
        set_cell_margins_zero(cell)
    return table


def _populate_columns_1_and_2(table, sides):
    row = table.rows[0]
    col1, col2, _ = row.cells

    col1_plan = _build_column_1_plan(sides)
    expected = _expected_row_count(sides)
    if len(col1_plan) != expected:
        raise ValueError(
            f"Column 1 plan has {len(col1_plan)} rows, expected {expected}"
        )

    clear_cell(col1)
    for text, indent in col1_plan:
        if indent is not None:
            add_para(col1, text, left_indent=Inches(indent))
        else:
            add_para(col1, text)

    clear_cell(col2)
    for _ in range(expected):
        add_para(col2, '§', alignment=WD_ALIGN_PARAGRAPH.CENTER)
    return expected


# ─── TX state caption builder ────────────────────────────────────────

def build_tx_state_caption(case):
    """Build a fresh document containing only the TX state caption."""
    doc = _common_doc_setup()

    add_para(doc, f'CAUSE NO. {case["cause_number"]}',
             alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, '')

    table = _build_caption_table(doc)
    sides = case["sides"]
    expected = _populate_columns_1_and_2(table, sides)

    col3 = table.rows[0].cells[2]
    clear_cell(col3)
    R = WD_ALIGN_PARAGRAPH.RIGHT
    first_v = _first_v_row_index(sides)
    for j in range(expected):
        if j == 0:
            add_para(col3, f'IN THE {case["court_type_line"]}', alignment=R)
        elif j == first_v:
            add_para(col3, case["district"], alignment=R)
        elif j == expected - 1:
            add_para(col3, case["county_line"], alignment=R)
        else:
            add_para(col3, '', alignment=R)

    return doc


# ─── Federal caption builder ─────────────────────────────────────────

def build_federal_caption(case):
    """Build a fresh document containing only the federal caption."""
    doc = _common_doc_setup()

    header_lines = [
        'IN THE UNITED STATES DISTRICT COURT',
        f'FOR THE {case["district"].upper()}',
    ]
    if case.get('division'):
        header_lines.append(case['division'].upper())
    add_court_header(doc, header_lines)

    table = _build_caption_table(doc)
    sides = case["sides"]
    expected = _populate_columns_1_and_2(table, sides)

    col3 = table.rows[0].cells[2]
    clear_cell(col3)
    R = WD_ALIGN_PARAGRAPH.RIGHT
    first_v = _first_v_row_index(sides)
    prefix = case.get('case_number_prefix', 'Civil Action No.')
    case_no = case['case_number']
    for j in range(expected):
        if j == first_v:
            add_para(col3, f'{prefix} {case_no}', alignment=R)
        else:
            add_para(col3, '', alignment=R)

    return doc


# ─── AAA arbitration caption builder ────────────────────────────────

def build_aaa_arb_caption(case):
    """Build a fresh document containing only the AAA arbitration caption."""
    doc = _common_doc_setup()

    # Forum header -- centered, plain weight (NOT bold), spaceAfter=12pt
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(12)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run('BEFORE THE AMERICAN ARBITRATION ASSOCIATION')
    set_run_font(run, bold=False)

    table = _build_caption_table(doc)
    sides = case["sides"]
    expected = _populate_columns_1_and_2(table, sides)

    col3 = table.rows[0].cells[2]
    clear_cell(col3)
    R = WD_ALIGN_PARAGRAPH.RIGHT
    first_v = _first_v_row_index(sides)
    case_no = case['case_number']
    for j in range(expected):
        if j == first_v:
            add_para(col3, f'AAA Case No. {case_no}', alignment=R)
        else:
            add_para(col3, '', alignment=R)

    return doc


# ─── Dispatcher ──────────────────────────────────────────────────────

BUILDERS = {
    'tx-state': build_tx_state_caption,
    'federal': build_federal_caption,
    'aaa-arb': build_aaa_arb_caption,
}


def build_caption(case, court_type='tx-state'):
    if court_type not in BUILDERS:
        raise NotImplementedError(
            f"Court type {court_type!r} not implemented. "
            f"Available: {sorted(BUILDERS)}"
        )
    return BUILDERS[court_type](case)



# ─── Splice ──────────────────────────────────────────────────────────────────

def _find_tx_state_caption_start(body_children):
    """TX state: caption starts at the 'CAUSE NO ...' paragraph."""
    for idx, child in enumerate(body_children):
        if child.tag == qn('w:p'):
            text_runs = child.findall('.//' + qn('w:t'))
            text = ''.join((t.text or '') for t in text_runs).strip().upper()
            if text.startswith('CAUSE NO'):
                return idx
    return None


def _find_federal_caption_start(body_children):
    """Federal: caption starts at the 'IN THE UNITED STATES DISTRICT COURT' paragraph."""
    for idx, child in enumerate(body_children):
        if child.tag == qn('w:p'):
            text_runs = child.findall('.//' + qn('w:t'))
            text = ''.join((t.text or '') for t in text_runs).strip().upper()
            if 'IN THE UNITED STATES DISTRICT COURT' in text:
                return idx
    return None


def _find_aaa_arb_caption_start(body_children):
    """AAA arb: caption starts at the 'BEFORE THE AMERICAN ARBITRATION ASSOCIATION' paragraph."""
    for idx, child in enumerate(body_children):
        if child.tag == qn('w:p'):
            text_runs = child.findall('.//' + qn('w:t'))
            text = ''.join((t.text or '') for t in text_runs).strip().upper()
            if 'BEFORE THE AMERICAN ARBITRATION' in text:
                return idx
    return None


CAPTION_FINDERS = {
    'tx-state': _find_tx_state_caption_start,
    'federal': _find_federal_caption_start,
    'aaa-arb': _find_aaa_arb_caption_start,
}


def _find_caption_end(body_children, start_idx):
    """Find the end of the caption block (first element after the caption table)."""
    found_table = False
    for idx in range(start_idx, len(body_children)):
        child = body_children[idx]
        if child.tag == qn('w:tbl'):
            found_table = True
        elif found_table:
            return idx
    return len(body_children)


def splice_caption_into(target_path, fresh_doc, court_type='tx-state'):
    """Replace the caption in target_path with the caption from fresh_doc."""
    from docx import Document as D
    from copy import deepcopy
    import zipfile, shutil, tempfile

    target = D(target_path)
    target_body = target.element.body
    target_children = list(target_body)

    finder = CAPTION_FINDERS.get(court_type)
    if not finder:
        raise NotImplementedError(f"No caption finder for {court_type!r}")

    start = finder(target_children)
    if start is None:
        raise ValueError(f"Could not find caption start in target document for {court_type}")

    end = _find_caption_end(target_children, start)

    fresh_body = fresh_doc.element.body
    fresh_children = list(fresh_body)

    # Remove old caption elements
    for child in target_children[start:end]:
        target_body.remove(child)

    # Insert fresh caption elements at the start position
    ref = target_children[end] if end < len(target_children) else None
    for child in fresh_children:
        if child.tag in (qn('w:p'), qn('w:tbl')):
            copied = deepcopy(child)
            if ref is not None:
                target_body.insert(list(target_body).index(ref), copied)
            else:
                target_body.append(copied)

    target.save(target_path)
    print(f"Caption spliced into {target_path}")
